import io

from docx import Document

from models.schemas import TranscriptLine


def _fmt_srt_time(seconds: float) -> str:
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3_600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1000)
    return f"{h:02}:{m:02}:{s:02},{ms:03}"


def _fmt_vtt_time(seconds: float) -> str:
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3_600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1000)
    return f"{h:02}:{m:02}:{s:02}.{ms:03}"


def _line_text(line: TranscriptLine, use_english: bool) -> str:
    if use_english and line.text_en:
        return line.text_en
    return line.text


def build_txt(lines: list[TranscriptLine], use_english: bool = False) -> str:
    return "\n".join(_line_text(l, use_english) for l in lines)


def build_srt(lines: list[TranscriptLine], use_english: bool = False) -> str:
    blocks = []
    for idx, line in enumerate(lines, start=1):
        blocks.append(
            f"{idx}\n{_fmt_srt_time(line.start)} --> {_fmt_srt_time(line.end)}\n"
            f"{_line_text(line, use_english)}\n"
        )
    return "\n".join(blocks)


def build_vtt(lines: list[TranscriptLine], use_english: bool = False) -> str:
    blocks = ["WEBVTT\n"]
    for line in lines:
        blocks.append(
            f"{_fmt_vtt_time(line.start)} --> {_fmt_vtt_time(line.end)}\n"
            f"{_line_text(line, use_english)}\n"
        )
    return "\n".join(blocks)


def build_docx(lines: list[TranscriptLine], use_english: bool = False) -> bytes:
    doc = Document()
    doc.add_heading("Transcript", level=1)
    for line in lines:
        ts = _fmt_vtt_time(line.start)[:-4]
        doc.add_paragraph(f"[{ts}] {_line_text(line, use_english)}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
